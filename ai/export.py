"""
Módulo de exportación de reportes de análisis de IA
"""

import os
import json
import logging
from datetime import datetime
from typing import Dict, Any, Optional
from pathlib import Path
import tempfile

# Importaciones para exportación
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.platypus import PageBreak, Image as RLImage
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    from bs4 import BeautifulSoup
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

logger = logging.getLogger(__name__)

class ReportExporter:
    """Exportador de reportes de análisis de IA"""
    
    def __init__(self, output_dir: str = "data/reports"):
        """
        Inicializa el exportador
        
        Args:
            output_dir: Directorio de salida para reportes
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Configuración de branding
        self.brand_config = {
            'clinic_name': 'Clínica Bonsana',
            'department': 'Departamento de TI',
            'primary_color': '#dc3545',
            'secondary_color': '#6c757d',
            'logo_path': None  # Ruta al logo si está disponible
        }
    
    def export_to_pdf(self, analysis_data: Dict[str, Any], filename: str = None) -> str:
        """
        Exporta análisis a PDF
        
        Args:
            analysis_data: Datos del análisis
            filename: Nombre del archivo (opcional)
            
        Returns:
            Ruta del archivo generado
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab no está disponible. Instalar con: pip install reportlab")
        
        if not filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            analysis_type = analysis_data.get('analysis_type', 'analysis')
            filename = f"report_{analysis_type}_{timestamp}.pdf"
        
        filepath = self.output_dir / filename
        
        try:
            # Crear documento PDF
            doc = SimpleDocTemplate(str(filepath), pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Personalizar estilos
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=HexColor(self.brand_config['primary_color']),
                alignment=1  # Centro
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceBefore=20,
                spaceAfter=12,
                textColor=HexColor(self.brand_config['primary_color'])
            )
            
            # Título principal
            story.append(Paragraph(f"Análisis de IA - {self.brand_config['clinic_name']}", title_style))
            story.append(Spacer(1, 20))
            
            # Información del reporte
            info_data = [
                ['Tipo de Análisis:', analysis_data.get('analysis_type', 'N/A')],
                ['Fecha de Generación:', datetime.now().strftime('%d/%m/%Y %H:%M')],
                ['Modelo de IA:', analysis_data.get('model_used', 'N/A')],
                ['Tiempo de Procesamiento:', f"{analysis_data.get('processing_time', 0):.2f} segundos"]
            ]
            
            info_table = Table(info_data, colWidths=[2*inch, 3*inch])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), HexColor('#f8f9fa')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(info_table)
            story.append(Spacer(1, 30))
            
            # Contenido del análisis
            analysis_text = analysis_data.get('analysis', '')
            
            if MARKDOWN_AVAILABLE:
                # Convertir markdown a HTML y luego procesar
                html_content = markdown.markdown(analysis_text)
                story.extend(self._html_to_pdf_elements(html_content, styles))
            else:
                # Procesar como texto plano
                story.extend(self._text_to_pdf_elements(analysis_text, styles))
            
            # Pie de página con metadatos
            story.append(PageBreak())
            story.append(Paragraph("Metadatos del Análisis", heading_style))
            
            context_data = analysis_data.get('context_used', {})
            if context_data:
                metadata_items = [
                    f"Total de tickets analizados: {context_data.get('total_tickets', 'N/A')}",
                    f"Tasa de resolución: {context_data.get('resolution_rate', 'N/A')}%",
                    f"Cumplimiento SLA: {context_data.get('sla_compliance', 'N/A')}%",
                    f"CSAT promedio: {context_data.get('average_csat', 'N/A')}"
                ]
                
                for item in metadata_items:
                    story.append(Paragraph(f"• {item}", styles['Normal']))
                    story.append(Spacer(1, 6))
            
            # Generar PDF
            doc.build(story)
            
            logger.info(f"Reporte PDF generado: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generando PDF: {str(e)}")
            raise
    
    def export_to_word(self, analysis_data: Dict[str, Any], filename: str = None) -> str:
        """
        Exporta análisis a Word
        
        Args:
            analysis_data: Datos del análisis
            filename: Nombre del archivo (opcional)
            
        Returns:
            Ruta del archivo generado
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está disponible. Instalar con: pip install python-docx")
        
        if not filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            analysis_type = analysis_data.get('analysis_type', 'analysis')
            filename = f"report_{analysis_type}_{timestamp}.docx"
        
        filepath = self.output_dir / filename
        
        try:
            # Crear documento Word
            doc = Document()
            
            # Configurar estilos
            styles = doc.styles
            
            # Estilo de título personalizado
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_format = title_style.paragraph_format
            title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Título principal
            title = doc.add_paragraph(f"Análisis de IA - {self.brand_config['clinic_name']}", style='CustomTitle')
            doc.add_paragraph()
            
            # Información del reporte
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Tipo de Análisis:', analysis_data.get('analysis_type', 'N/A')),
                ('Fecha de Generación:', datetime.now().strftime('%d/%m/%Y %H:%M')),
                ('Modelo de IA:', analysis_data.get('model_used', 'N/A')),
                ('Tiempo de Procesamiento:', f"{analysis_data.get('processing_time', 0):.2f} segundos")
            ]
            
            for i, (label, value) in enumerate(info_data):
                row = info_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(value)
                row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            doc.add_paragraph()
            
            # Contenido del análisis
            analysis_text = analysis_data.get('analysis', '')
            
            if MARKDOWN_AVAILABLE:
                # Convertir markdown a elementos de Word
                self._markdown_to_word(analysis_text, doc)
            else:
                # Agregar como texto plano
                doc.add_paragraph(analysis_text)
            
            # Guardar documento
            doc.save(str(filepath))
            
            logger.info(f"Reporte Word generado: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generando Word: {str(e)}")
            raise
    
    def export_to_json(self, analysis_data: Dict[str, Any], filename: str = None) -> str:
        """
        Exporta análisis a JSON estructurado
        
        Args:
            analysis_data: Datos del análisis
            filename: Nombre del archivo (opcional)
            
        Returns:
            Ruta del archivo generado
        """
        if not filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            analysis_type = analysis_data.get('analysis_type', 'analysis')
            filename = f"report_{analysis_type}_{timestamp}.json"
        
        filepath = self.output_dir / filename
        
        try:
            # Crear reporte estructurado
            structured_report = {
                'metadata': {
                    'clinic_name': self.brand_config['clinic_name'],
                    'department': self.brand_config['department'],
                    'export_timestamp': datetime.now().isoformat(),
                    'analysis_type': analysis_data.get('analysis_type', 'unknown'),
                    'model_used': analysis_data.get('model_used', 'unknown'),
                    'processing_time_seconds': analysis_data.get('processing_time', 0),
                    'success': analysis_data.get('success', False)
                },
                'analysis': {
                    'raw_text': analysis_data.get('analysis', ''),
                    'sections': self._extract_sections_from_text(analysis_data.get('analysis', '')),
                    'summary': self._generate_summary(analysis_data.get('analysis', ''))
                },
                'context': analysis_data.get('context_used', {}),
                'metrics': {
                    'word_count': len(analysis_data.get('analysis', '').split()),
                    'response_tokens': analysis_data.get('response_tokens', 0),
                    'prompt_tokens': analysis_data.get('prompt_tokens', 0)
                }
            }
            
            # Guardar JSON
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(structured_report, f, indent=2, ensure_ascii=False, default=str)
            
            logger.info(f"Reporte JSON generado: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generando JSON: {str(e)}")
            raise
    
    def export_to_html(self, analysis_data: Dict[str, Any], filename: str = None) -> str:
        """
        Exporta análisis a HTML
        
        Args:
            analysis_data: Datos del análisis
            filename: Nombre del archivo (opcional)
            
        Returns:
            Ruta del archivo generado
        """
        if not filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            analysis_type = analysis_data.get('analysis_type', 'analysis')
            filename = f"report_{analysis_type}_{timestamp}.html"
        
        filepath = self.output_dir / filename
        
        try:
            # Convertir markdown a HTML si está disponible
            if MARKDOWN_AVAILABLE:
                analysis_html = markdown.markdown(
                    analysis_data.get('analysis', ''),
                    extensions=['tables', 'fenced_code', 'toc']
                )
            else:
                # Convertir texto plano a HTML básico
                analysis_text = analysis_data.get('analysis', '')
                analysis_html = analysis_text.replace('\n\n', '</p><p>').replace('\n', '<br>')
                analysis_html = f"<p>{analysis_html}</p>"
            
            # Template HTML completo
            html_template = f"""
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Análisis de IA - {self.brand_config['clinic_name']}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f8f9fa;
        }}
        .header {{
            background: linear-gradient(135deg, {self.brand_config['primary_color']} 0%, #b02a37 100%);
            color: white;
            padding: 2rem;
            border-radius: 12px;
            margin-bottom: 2rem;
            text-align: center;
        }}
        .header h1 {{
            margin: 0;
            font-size: 2.5rem;
        }}
        .header p {{
            margin: 0.5rem 0 0 0;
            opacity: 0.9;
        }}
        .content {{
            background: white;
            padding: 2rem;
            border-radius: 12px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            margin-bottom: 2rem;
        }}
        .metadata {{
            background: #e9ecef;
            padding: 1rem;
            border-radius: 8px;
            margin-bottom: 2rem;
        }}
        .metadata table {{
            width: 100%;
            border-collapse: collapse;
        }}
        .metadata td {{
            padding: 0.5rem;
            border-bottom: 1px solid #dee2e6;
        }}
        .metadata td:first-child {{
            font-weight: bold;
            width: 200px;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: {self.brand_config['primary_color']};
            margin-top: 2rem;
            margin-bottom: 1rem;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 1rem 0;
        }}
        th, td {{
            border: 1px solid #dee2e6;
            padding: 0.75rem;
            text-align: left;
        }}
        th {{
            background-color: {self.brand_config['primary_color']};
            color: white;
        }}
        code {{
            background-color: #f8f9fa;
            padding: 0.2rem 0.4rem;
            border-radius: 4px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #f8f9fa;
            padding: 1rem;
            border-radius: 8px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid {self.brand_config['primary_color']};
            background-color: #f8d7da;
            padding: 1rem;
            margin: 1rem 0;
            border-radius: 0 8px 8px 0;
        }}
        .footer {{
            text-align: center;
            color: {self.brand_config['secondary_color']};
            margin-top: 2rem;
            padding: 1rem;
            border-top: 1px solid #dee2e6;
        }}
        @media print {{
            body {{ background: white; }}
            .header {{ background: {self.brand_config['primary_color']}; }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Análisis de IA</h1>
        <p>{self.brand_config['clinic_name']} - {self.brand_config['department']}</p>
    </div>
    
    <div class="content">
        <div class="metadata">
            <table>
                <tr>
                    <td>Tipo de Análisis:</td>
                    <td>{analysis_data.get('analysis_type', 'N/A')}</td>
                </tr>
                <tr>
                    <td>Fecha de Generación:</td>
                    <td>{datetime.now().strftime('%d/%m/%Y %H:%M')}</td>
                </tr>
                <tr>
                    <td>Modelo de IA:</td>
                    <td>{analysis_data.get('model_used', 'N/A')}</td>
                </tr>
                <tr>
                    <td>Tiempo de Procesamiento:</td>
                    <td>{analysis_data.get('processing_time', 0):.2f} segundos</td>
                </tr>
            </table>
        </div>
        
        <div class="analysis-content">
            {analysis_html}
        </div>
    </div>
    
    <div class="footer">
        <p>Generado por el Sistema de Análisis de IA - {self.brand_config['clinic_name']}</p>
        <p>Este reporte es confidencial y está destinado únicamente para uso interno.</p>
    </div>
</body>
</html>
"""
            
            # Guardar HTML
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_template)
            
            logger.info(f"Reporte HTML generado: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Error generando HTML: {str(e)}")
            raise
    
    def _html_to_pdf_elements(self, html_content: str, styles):
        """Convierte HTML a elementos PDF"""
        elements = []
        
        if MARKDOWN_AVAILABLE:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
                if element.name.startswith('h'):
                    level = int(element.name[1])
                    style_name = f'Heading{min(level, 3)}'
                    elements.append(Paragraph(element.get_text(), styles[style_name]))
                    elements.append(Spacer(1, 12))
                elif element.name == 'p':
                    elements.append(Paragraph(element.get_text(), styles['Normal']))
                    elements.append(Spacer(1, 6))
        
        return elements
    
    def _text_to_pdf_elements(self, text: str, styles):
        """Convierte texto plano a elementos PDF"""
        elements = []
        
        for paragraph in text.split('\n\n'):
            if paragraph.strip():
                elements.append(Paragraph(paragraph.strip(), styles['Normal']))
                elements.append(Spacer(1, 12))
        
        return elements
    
    def _markdown_to_word(self, markdown_text: str, doc):
        """Convierte markdown a elementos de Word"""
        if not MARKDOWN_AVAILABLE:
            doc.add_paragraph(markdown_text)
            return
        
        html = markdown.markdown(markdown_text)
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                heading_style = f'Heading {min(level, 9)}'
                doc.add_paragraph(element.get_text(), style=heading_style)
            elif element.name == 'p':
                doc.add_paragraph(element.get_text())
    
    def _extract_sections_from_text(self, text: str) -> Dict[str, str]:
        """Extrae secciones del texto basado en headers"""
        sections = {}
        current_section = "introduccion"
        current_content = []
        
        lines = text.split('\n')
        
        for line in lines:
            if line.strip().startswith('#'):
                # Guardar sección anterior
                if current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                # Nueva sección
                header_text = line.strip().lstrip('#').strip()
                current_section = header_text.lower().replace(' ', '_').replace(':', '')
                current_content = []
            else:
                current_content.append(line)
        
        # Guardar última sección
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    def _generate_summary(self, text: str, max_sentences: int = 3) -> str:
        """Genera resumen del texto"""
        sentences = text.split('.')[:max_sentences]
        return '. '.join(sentence.strip() for sentence in sentences if sentence.strip()) + '.'
    
    def get_available_formats(self) -> Dict[str, bool]:
        """
        Obtiene formatos de exportación disponibles
        
        Returns:
            Diccionario con formatos disponibles
        """
        return {
            'pdf': REPORTLAB_AVAILABLE,
            'word': DOCX_AVAILABLE,
            'html': True,
            'json': True
        }
    
    def export_all_formats(self, analysis_data: Dict[str, Any], base_filename: str = None) -> Dict[str, str]:
        """
        Exporta análisis en todos los formatos disponibles
        
        Args:
            analysis_data: Datos del análisis
            base_filename: Nombre base del archivo (sin extensión)
            
        Returns:
            Diccionario con rutas de archivos generados
        """
        if not base_filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            analysis_type = analysis_data.get('analysis_type', 'analysis')
            base_filename = f"report_{analysis_type}_{timestamp}"
        
        exported_files = {}
        
        # JSON (siempre disponible)
        try:
            json_file = self.export_to_json(analysis_data, f"{base_filename}.json")
            exported_files['json'] = json_file
        except Exception as e:
            logger.error(f"Error exportando JSON: {e}")
        
        # HTML (siempre disponible)
        try:
            html_file = self.export_to_html(analysis_data, f"{base_filename}.html")
            exported_files['html'] = html_file
        except Exception as e:
            logger.error(f"Error exportando HTML: {e}")
        
        # PDF (si está disponible)
        if REPORTLAB_AVAILABLE:
            try:
                pdf_file = self.export_to_pdf(analysis_data, f"{base_filename}.pdf")
                exported_files['pdf'] = pdf_file
            except Exception as e:
                logger.error(f"Error exportando PDF: {e}")
        
        # Word (si está disponible)
        if DOCX_AVAILABLE:
            try:
                word_file = self.export_to_word(analysis_data, f"{base_filename}.docx")
                exported_files['word'] = word_file
            except Exception as e:
                logger.error(f"Error exportando Word: {e}")
        
        logger.info(f"Exportación completada: {len(exported_files)} formatos generados")
        return exported_files